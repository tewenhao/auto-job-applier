"""Tests for the deterministic ingestion pieces (no LLM / network)."""

from __future__ import annotations

from datetime import date
from pathlib import Path
from uuid import uuid4

import pytest

from app.ingestion import pipeline as pl
from app.ingestion.documents import extract_text, iter_documents
from app.ingestion.github import build_github_profile
from app.ingestion.mapping import apply_contact, normalize_date, to_experience, to_skill
from app.ingestion.pipeline import Ingestor
from app.ingestion.schema import (
    ExtractedContact,
    ExtractedExperience,
    ExtractedSkill,
    ProfileExtraction,
)
from app.profile.models import Candidate, ExperienceKind, SourceType

CID = uuid4()


class _FakeRepo:
    """Records how experiences were written (upsert vs plain insert)."""

    def __init__(self) -> None:
        self.added: list[object] = []
        self.upserted: list[object] = []

    def add_source_document(self, doc):  # type: ignore[no-untyped-def]
        doc.id = uuid4()
        return doc

    def get_candidate(self, candidate_id):  # type: ignore[no-untyped-def]
        return None

    def upsert_experience(self, exp):  # type: ignore[no-untyped-def]
        self.upserted.append(exp)
        return exp

    def add_experience(self, exp):  # type: ignore[no-untyped-def]
        self.added.append(exp)
        return exp

    def upsert_skill(self, skill):  # type: ignore[no-untyped-def]
        return skill


def test_fresh_ingest_inserts_without_dedup(tmp_path: Path, monkeypatch) -> None:  # type: ignore[no-untyped-def]
    # Two distinct entries that share the natural key (same kind, null org/title):
    # with dedup they'd collapse via upsert; after --fresh they must both survive.
    twins = [
        ExtractedExperience(kind=ExperienceKind.PROJECT, summary="Hackathon A"),
        ExtractedExperience(kind=ExperienceKind.PROJECT, summary="Hackathon B"),
    ]
    monkeypatch.setattr(pl, "extract_profile", lambda *a, **k: ProfileExtraction(experiences=twins))
    md = tmp_path / "master.md"
    md.write_text("content", encoding="utf-8")

    fresh_repo = _FakeRepo()
    summary = Ingestor(fresh_repo, llm=None).ingest_document(  # type: ignore[arg-type]
        md, SourceType.MASTER_DOC, candidate_id=CID, dedup=False
    )
    assert summary["experiences"] == 2
    assert len(fresh_repo.added) == 2 and len(fresh_repo.upserted) == 0

    incremental_repo = _FakeRepo()
    Ingestor(incremental_repo, llm=None).ingest_document(  # type: ignore[arg-type]
        md, SourceType.MASTER_DOC, candidate_id=CID, dedup=True
    )
    assert len(incremental_repo.upserted) == 2 and len(incremental_repo.added) == 0


# --- documents ---
def test_extract_text_txt_and_md(tmp_path: Path) -> None:
    txt = tmp_path / "a.txt"
    txt.write_text("hello world", encoding="utf-8")
    assert extract_text(txt) == "hello world"

    md = tmp_path / "b.md"
    md.write_text("# Title\n\nBody", encoding="utf-8")
    assert "Body" in extract_text(md)


def test_extract_text_docx(tmp_path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("First line")
    doc.add_paragraph("Second line")
    p = tmp_path / "c.docx"
    doc.save(str(p))
    text = extract_text(p)
    assert "First line" in text and "Second line" in text


def test_extract_text_unsupported(tmp_path: Path) -> None:
    p = tmp_path / "d.rtf"
    p.write_text("x", encoding="utf-8")
    with pytest.raises(ValueError):
        extract_text(p)


def test_extract_text_missing() -> None:
    with pytest.raises(FileNotFoundError):
        extract_text("/nonexistent/file.txt")


def test_iter_documents_single_file(tmp_path: Path) -> None:
    f = tmp_path / "one.txt"
    f.write_text("x", encoding="utf-8")
    assert iter_documents(f) == [f]


def test_iter_documents_directory_filters_and_sorts(tmp_path: Path) -> None:
    (tmp_path / "b.pdf").write_text("x", encoding="utf-8")
    (tmp_path / "a.docx").write_text("x", encoding="utf-8")
    (tmp_path / "c.txt").write_text("x", encoding="utf-8")
    (tmp_path / "notes.rtf").write_text("x", encoding="utf-8")  # unsupported
    (tmp_path / ".hidden.pdf").write_text("x", encoding="utf-8")  # hidden
    nested = tmp_path / "old"
    nested.mkdir()
    (nested / "d.md").write_text("x", encoding="utf-8")  # recursive

    names = [p.name for p in iter_documents(tmp_path)]
    assert names == ["a.docx", "b.pdf", "c.txt", "d.md"]


def test_iter_documents_nonexistent_returns_path() -> None:
    # A bare (missing) path is returned as-is; extract_text raises on use.
    assert iter_documents("/nope/x.pdf") == [Path("/nope/x.pdf")]


# --- mapping ---
def test_normalize_date_variants() -> None:
    assert normalize_date("2025-06-01") == date(2025, 6, 1)
    assert normalize_date("2025-06") == date(2025, 6, 1)
    assert normalize_date("2025") == date(2025, 1, 1)
    assert normalize_date("") is None
    assert normalize_date(None) is None
    assert normalize_date("Summer 2025") is None


def test_to_experience_attaches_identity_and_dates() -> None:
    x = ExtractedExperience(
        kind=ExperienceKind.WORK, org="Acme", title="Intern", start="2025-06", is_current=True
    )
    doc_id = uuid4()
    exp = to_experience(x, candidate_id=CID, source="resume", source_document_id=doc_id)
    assert exp.candidate_id == CID
    assert exp.source == "resume"
    assert exp.source_document_id == doc_id
    assert exp.start_date == date(2025, 6, 1)
    assert exp.is_current is True


def test_to_skill() -> None:
    s = to_skill(ExtractedSkill(name="Python", category="Languages"), candidate_id=CID)
    assert s.name == "Python" and s.candidate_id == CID


def test_apply_contact_fills_only_empty_fields() -> None:
    candidate = Candidate(id=CID, full_name="Existing Name")
    contact = ExtractedContact(full_name="New Name", email="new@example.com")
    updated = apply_contact(candidate, contact)
    assert updated.full_name == "Existing Name"  # not clobbered
    assert updated.email == "new@example.com"  # filled in


# --- github ---
def test_build_github_profile_aggregates() -> None:
    user = {"login": "jlee", "public_repos": 2, "followers": 10, "name": "J Lee"}
    repos = [
        {"name": "a", "language": "Python", "stargazers_count": 5, "topics": ["ml"]},
        {"name": "b", "language": "Python", "fork": True},
        {"name": "c", "language": "Rust"},
    ]
    gh = build_github_profile(user, repos, candidate_id=CID)
    assert gh.username == "jlee"
    assert gh.languages == {"Python": 2, "Rust": 1}  # most_common order
    assert len(gh.repos) == 3
    assert gh.repos[0]["stars"] == 5
    assert gh.stats["followers"] == 10
    assert gh.pulled_at is not None
